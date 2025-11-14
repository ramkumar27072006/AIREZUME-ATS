# backend/template_engine.py
import os
import tempfile
from typing import Dict, Any
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import uuid

def generate_docx(enhanced: Dict[str, Any]) -> str:
    """
    Generate a .docx from enhanced dict and return path.
    """
    doc = Document()
    name = enhanced.get("name", "")
    summary = enhanced.get("summary", "")
    bullets = enhanced.get("bullets", []) or []
    text = enhanced.get("text", "")

    if name:
        doc.add_heading(name, level=1)
    if summary:
        doc.add_paragraph(summary)
    if bullets:
        doc.add_paragraph("")
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet')

    # Also append raw text at end
    if text:
        doc.add_paragraph("")
        doc.add_paragraph(text)

    tmpdir = tempfile.gettempdir()
    fname = f"enhanced_{uuid.uuid4().hex[:8]}.docx"
    path = os.path.join(tmpdir, fname)
    doc.save(path)
    return path

def generate_pdf_from_text(text: str, name: str = "candidate") -> str:
    """
    Simple PDF generation using reportlab. Returns file path.
    """
    tmpdir = tempfile.gettempdir()
    fname = f"resume_{uuid.uuid4().hex[:8]}.pdf"
    path = os.path.join(tmpdir, fname)

    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    margin = 50
    max_width = width - 2 * margin
    y = height - margin

    # Title
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, name)
    y -= 24

    c.setFont("Helvetica", 10)
    lines = text.splitlines()
    for line in lines:
        # wrap long lines
        while line and y > margin:
            if len(line) < 90:
                c.drawString(margin, y, line)
                y -= 14
                break
            else:
                # take slice approx
                part = line[:90]
                # try to cut at space
                last_space = part.rfind(" ")
                if last_space > 0:
                    part = line[:last_space]
                c.drawString(margin, y, part)
                y -= 14
                line = line[len(part):].lstrip()
        if y <= margin + 20:
            c.showPage()
            y = height - margin
            c.setFont("Helvetica", 10)
    c.save()
    return path
